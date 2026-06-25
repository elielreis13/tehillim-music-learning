"""
Script para gerar o documento de QA da plataforma Tehillim Music Learning.
Execute: python3 generate_qa_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_test_table(doc, test_id, test_name, objective, preconditions, steps, expected, module_prefix=""):
    """Adiciona uma tabela de caso de teste completa ao documento."""
    doc.add_paragraph()  # Espaço antes

    # Título do teste
    p = doc.add_paragraph()
    run = p.add_run(f"  {test_id} — {test_name}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(6)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Definir larguras das colunas
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(13)

    def add_row(label, content, header=False, label_bg="E8C77A", content_bg=None):
        row = table.add_row()
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)

        cell_label = row.cells[0]
        cell_content = row.cells[1]

        set_cell_bg(cell_label, label_bg)
        if content_bg:
            set_cell_bg(cell_content, content_bg)

        p_label = cell_label.paragraphs[0]
        run_label = p_label.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = RGBColor(0x2A, 0x24, 0x19)

        if isinstance(content, list):
            for i, item in enumerate(content):
                if i == 0:
                    p_c = cell_content.paragraphs[0]
                else:
                    p_c = cell_content.add_paragraph()
                p_c.add_run(item).font.size = Pt(9)
        else:
            p_c = cell_content.paragraphs[0]
            p_c.add_run(content).font.size = Pt(9)

        return row

    add_row("ID do Teste", test_id, label_bg="C4943A")
    add_row("Nome do Teste", test_name)
    add_row("Objetivo", objective)

    if isinstance(preconditions, list):
        pre_list = preconditions
    else:
        pre_list = [preconditions]
    add_row("Pré-condições", pre_list)

    steps_list = [f"{i+1}. {s}" for i, s in enumerate(steps)]
    add_row("Passo a Passo", steps_list)

    if isinstance(expected, list):
        exp_list = expected
    else:
        exp_list = [expected]
    add_row("Resultado Esperado", exp_list)

    add_row("Resultado Obtido", "_______________________________________________", label_bg="D4EDDA", content_bg="F8FFF9")
    add_row("Status", "[ ] Aprovado     [ ] Reprovado     [ ] Bloqueado", label_bg="FFF3CD", content_bg="FFFEF5")
    add_row("Evidências", "Anexar print(s) ou descrever o comportamento observado:", label_bg="D1ECF1", content_bg="F5FBFC")
    add_row("Observações", "_______________________________________________", label_bg="F8D7DA", content_bg="FFFAFA")

    doc.add_paragraph()  # Espaço depois


# ─────────────────────────────────────────────
# DOCUMENTO PRINCIPAL
# ─────────────────────────────────────────────

doc = Document()

# Configurar margens
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Estilo padrão
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ─── CAPA ───────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TEHILLIM MUSIC LEARNING")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0xC4, 0x94, 0x3A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plataforma de Ensino Musical Online")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4E, 0x7A, 0x60)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DOCUMENTO DE QUALIDADE — QA")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Checklist Oficial de Homologação — Release v1.0")
run.font.size = Pt(12)

doc.add_paragraph("\n\n")

# Tabela de informações iniciais
info_table = doc.add_table(rows=0, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

def add_info_row(table, label, value):
    row = table.add_row()
    set_cell_bg(row.cells[0], "E8C77A")
    r0 = row.cells[0].paragraphs[0].add_run(label)
    r0.bold = True
    r0.font.size = Pt(10)
    r1 = row.cells[1].paragraphs[0].add_run(value)
    r1.font.size = Pt(10)

add_info_row(info_table, "Nome da Aplicação", "Tehillim Music Learning")
add_info_row(info_table, "Versão do Release", "v1.0 — Release Inicial")
add_info_row(info_table, "Data da Validação", "____/____/________")
add_info_row(info_table, "Responsável pelo QA", "___________________________________")
add_info_row(info_table, "Ambiente Testado", "[ ] DEV   [ ] STAGING   [ ] PRODUÇÃO")
add_info_row(info_table, "Navegador Utilizado", "[ ] Chrome   [ ] Firefox   [ ] Safari   [ ] Edge")
add_info_row(info_table, "Sistema Operacional", "[ ] Windows   [ ] macOS   [ ] Linux   [ ] iOS   [ ] Android")
add_info_row(info_table, "Observações Gerais", "_______________________________________________")

doc.add_page_break()

# ─── SUMÁRIO MANUAL ─────────────────────────
add_heading(doc, "SUMÁRIO", level=1, color=(0x1A, 0x1A, 0x2E))

modulos_sumario = [
    ("1.", "Módulo de Autenticação (AUTH)"),
    ("2.", "Módulo de Navegação e Rotas (NAV)"),
    ("3.", "Módulo Dashboard do Aluno (DASH)"),
    ("4.", "Módulo de Grupos de Conteúdo (GRP)"),
    ("5.", "Módulo de Conteúdo e Steps (MOD)"),
    ("6.", "Módulo de Exercícios (EXR)"),
    ("7.", "Módulo de Jogos Interativos (GAME)"),
    ("8.", "Módulo Player Bona — Partitura (BONA)"),
    ("9.", "Módulo de Trilhas de Aprendizado (TRAIL)"),
    ("10.", "Módulo de Desempenho (PERF)"),
    ("11.", "Módulo de Conquistas (ACH)"),
    ("12.", "Módulo de Mensagens — Aluno (MSG)"),
    ("13.", "Módulo de Notificações (NOTIF)"),
    ("14.", "Módulo de Configurações e Perfil (CONF)"),
    ("15.", "Módulo de Gravação de Áudio (REC)"),
    ("16.", "Módulo Dashboard do Professor (TDASH)"),
    ("17.", "Módulo Painel Admin do Professor (TADM)"),
    ("18.", "Módulo Gestão de Alunos pelo Professor (TSTU)"),
    ("19.", "Testes de Responsividade (RESP)"),
    ("20.", "Testes de Segurança Funcional (SEC)"),
    ("21.", "Testes de Performance Visual (VIS)"),
    ("22.", "Testes de Estados Extremos (EXTR)"),
    ("23.", "Testes de Acessibilidade (ACESS)"),
]

for num, titulo in modulos_sumario:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{num}  {titulo}")
    run.font.size = Pt(11)

doc.add_page_break()

# ─── SEÇÃO: CONTEXTO DA APLICAÇÃO ──────────
add_heading(doc, "CONTEXTO DA APLICAÇÃO", level=1, color=(0x4E, 0x7A, 0x60))

context_text = """A Tehillim Music Learning é uma plataforma de ensino musical online voltada para alunos que desejam aprender teoria musical, solfejo, ritmo e harmonia de forma interativa e progressiva.

A plataforma é composta por dois perfis principais de usuário:

• ALUNO (padrão): acessa módulos liberados pelo professor, realiza exercícios, acompanha seu progresso, participa de jogos interativos, grava áudios e troca mensagens com o professor.

• PROFESSOR: gerencia os alunos, libera acesso a módulos, acompanha o progresso de cada aluno, adiciona aulas extras, deixa comentários e responde mensagens.

A aplicação é acessada via navegador web (desktop e mobile) e não requer instalação. O backend utiliza Flask (Python) com banco de dados Supabase (PostgreSQL) e autenticação via Supabase Auth.

GRUPOS DE CONTEÚDO DISPONÍVEIS:
1. Fundamentos do Som — Módulos 1 a 15
2. Musicalização e Iniciação — Módulos 16 a 30
3. Alfabetização Musical — Módulos 31 a 42
4. Ritmo — Módulos 43 a 60
5. Melodia e Solfejo — Módulos 61 a 69
6. Intervalos e Escalas — Módulos 70 a 78
7. Harmonia e Acordes — Módulos 79 a 90
8. Tonalidade e Análise — Módulos 91 a 93
9. Dinâmica e Expressão — Módulos 94 a 99 e 141 a 143
10. Método Bona (Ritmo Sistemático) — Módulos 101 a 140
11. Método Pozzoli (Solfejo Progressivo) — Módulos 201 a 212"""

p = doc.add_paragraph(context_text)
p.runs[0].font.size = Pt(10)

doc.add_page_break()


# ══════════════════════════════════════════════
# MÓDULO 1 — AUTENTICAÇÃO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 1 — AUTENTICAÇÃO (AUTH)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida todo o fluxo de entrada e saída da plataforma, controle de sessão, redirecionamentos e controle de acesso por papel (role).", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="AUTH-001",
    test_name="Login com credenciais válidas — Aluno",
    objective="Verificar que um aluno cadastrado consegue fazer login com e-mail e senha válidos e é redirecionado corretamente para a home.",
    preconditions=["Usuário aluno cadastrado e ativo no sistema.", "Nenhuma sessão ativa no navegador."],
    steps=[
        "Acessar a URL da aplicação (ex: https://app.tehillim.com.br).",
        "Verificar que a página de login é exibida.",
        "Inserir o e-mail de um aluno cadastrado no campo 'E-mail'.",
        "Inserir a senha correta no campo 'Senha'.",
        "Clicar no botão 'Entrar' (ou equivalente).",
        "Aguardar o processamento da autenticação.",
    ],
    expected=[
        "O sistema autentica o usuário com sucesso.",
        "O usuário é redirecionado para a página inicial do aluno (/inicio ou /).",
        "O nome do aluno aparece no header ou sidebar.",
        "O menu lateral (sidebar) exibe as opções: Início, Trilhas, Aulas, Desempenho, Conquistas, Mensagens, Notificações, Configurações.",
        "Nenhuma mensagem de erro é exibida.",
        "A URL muda para a home do aluno.",
    ]
)

add_test_table(doc,
    test_id="AUTH-002",
    test_name="Login com credenciais válidas — Professor",
    objective="Verificar que um professor consegue fazer login e é redirecionado para o dashboard do professor.",
    preconditions=["Usuário com perfil 'professor' cadastrado no sistema.", "Nenhuma sessão ativa no navegador."],
    steps=[
        "Acessar a URL de login da aplicação.",
        "Inserir o e-mail de um professor cadastrado.",
        "Inserir a senha correta.",
        "Clicar em 'Entrar'.",
        "Aguardar o processamento.",
    ],
    expected=[
        "O professor é autenticado com sucesso.",
        "O usuário é redirecionado para a home ou dashboard.",
        "O menu lateral exibe as opções adicionais de professor: 'Admin' e 'Professor/Alunos'.",
        "Os menus exclusivos de aluno (ex: Desempenho, Conquistas) podem estar visíveis ou ocultos conforme o design.",
        "Nenhum erro é exibido.",
    ]
)

add_test_table(doc,
    test_id="AUTH-003",
    test_name="Login com e-mail inválido / não cadastrado",
    objective="Verificar que o sistema rejeita credenciais de e-mail que não existem na base.",
    preconditions=["Nenhuma sessão ativa no navegador."],
    steps=[
        "Acessar a página de login.",
        "Inserir um e-mail que não existe na base (ex: naoexiste@teste.com).",
        "Inserir qualquer senha no campo 'Senha'.",
        "Clicar em 'Entrar'.",
    ],
    expected=[
        "O sistema NÃO autentica o usuário.",
        "Uma mensagem de erro clara é exibida (ex: 'E-mail ou senha inválidos', 'Usuário não encontrado').",
        "O usuário permanece na página de login.",
        "Nenhum dado sensível é exposto na mensagem de erro.",
        "A URL não muda para a home.",
    ]
)

add_test_table(doc,
    test_id="AUTH-004",
    test_name="Login com senha incorreta",
    objective="Verificar que o sistema rejeita login com senha errada para um e-mail válido.",
    preconditions=["Usuário cadastrado no sistema.", "Nenhuma sessão ativa."],
    steps=[
        "Acessar a página de login.",
        "Inserir o e-mail de um usuário cadastrado.",
        "Inserir uma senha incorreta.",
        "Clicar em 'Entrar'.",
    ],
    expected=[
        "O sistema NÃO autentica o usuário.",
        "Mensagem de erro é exibida (ex: 'E-mail ou senha inválidos').",
        "O usuário permanece na página de login.",
        "A senha incorreta NÃO é exibida em nenhum lugar da tela.",
    ]
)

add_test_table(doc,
    test_id="AUTH-005",
    test_name="Login com campos vazios",
    objective="Verificar que o formulário valida campos obrigatórios antes de submeter.",
    preconditions=["Nenhuma sessão ativa."],
    steps=[
        "Acessar a página de login.",
        "Deixar o campo 'E-mail' em branco.",
        "Deixar o campo 'Senha' em branco.",
        "Clicar em 'Entrar'.",
    ],
    expected=[
        "O formulário NÃO submete a requisição.",
        "Uma validação visual é exibida nos campos obrigatórios (ex: borda vermelha, mensagem 'Campo obrigatório').",
        "O usuário permanece na página de login.",
    ]
)

add_test_table(doc,
    test_id="AUTH-006",
    test_name="Login com apenas e-mail preenchido",
    objective="Verificar validação com senha vazia.",
    preconditions=["Nenhuma sessão ativa."],
    steps=[
        "Acessar a página de login.",
        "Inserir um e-mail válido.",
        "Deixar o campo 'Senha' em branco.",
        "Clicar em 'Entrar'.",
    ],
    expected=[
        "O formulário NÃO submete ou exibe erro.",
        "Indicação visual de que a senha é obrigatória.",
    ]
)

add_test_table(doc,
    test_id="AUTH-007",
    test_name="Logout do sistema",
    objective="Verificar que o usuário consegue encerrar a sessão corretamente.",
    preconditions=["Usuário logado no sistema."],
    steps=[
        "Localizar o botão ou opção de 'Sair' / 'Logout' na interface (geralmente no menu de configurações ou perfil).",
        "Clicar em 'Sair'.",
        "Aguardar o processamento.",
        "Tentar acessar qualquer página protegida (ex: /inicio).",
    ],
    expected=[
        "A sessão é encerrada com sucesso.",
        "O usuário é redirecionado para a página de login ou landing page.",
        "Ao tentar acessar /inicio sem sessão, o usuário é redirecionado para login.",
        "Nenhum dado do usuário anterior persiste na interface.",
    ]
)

add_test_table(doc,
    test_id="AUTH-008",
    test_name="Acesso à rota protegida sem login",
    objective="Verificar que rotas protegidas redirecionam usuários não autenticados para o login.",
    preconditions=["Nenhuma sessão ativa no navegador."],
    steps=[
        "Sem estar logado, digitar manualmente na barra de endereço: /inicio",
        "Pressionar Enter.",
        "Repetir para: /trilhas, /aulas, /desempenho, /conquistas, /mensagens, /configuracoes, /perfil",
    ],
    expected=[
        "Em TODAS as rotas listadas, o usuário é redirecionado para a página de login.",
        "Nenhuma página protegida é exibida sem autenticação.",
        "A URL muda para /login ou /.",
    ]
)

add_test_table(doc,
    test_id="AUTH-009",
    test_name="Acesso a rotas de professor por aluno",
    objective="Verificar que alunos não conseguem acessar páginas exclusivas de professores.",
    preconditions=["Usuário logado como ALUNO (não professor)."],
    steps=[
        "Estando logado como aluno, digitar manualmente: /dashboard",
        "Pressionar Enter.",
        "Repetir para: /admin, /professor/alunos",
    ],
    expected=[
        "O aluno é redirecionado para a home ou recebe mensagem de acesso negado.",
        "Nenhuma página de professor é exibida para o aluno.",
    ]
)

add_test_table(doc,
    test_id="AUTH-010",
    test_name="Persistência de sessão após fechar e reabrir o navegador",
    objective="Verificar que a sessão persiste conforme configuração (sessão de 7 dias).",
    preconditions=["Usuário logado com sucesso."],
    steps=[
        "Realizar login com sucesso.",
        "Fechar completamente o navegador.",
        "Reabrir o navegador.",
        "Acessar a URL da aplicação.",
    ],
    expected=[
        "O usuário ainda está logado (sessão persistida).",
        "A home do usuário é exibida sem necessidade de novo login.",
        "OU, se a sessão expirou, o usuário é redirecionado ao login — comportamento deve ser consistente.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 2 — NAVEGAÇÃO E ROTAS
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 2 — NAVEGAÇÃO E ROTAS (NAV)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a navegação entre páginas, menus, links, redirecionamentos e tratamento de rotas inválidas.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="NAV-001",
    test_name="Landing Page — Carregamento e conteúdo",
    objective="Verificar que a página pública (landing page) carrega corretamente para visitantes não logados.",
    preconditions=["Nenhuma sessão ativa."],
    steps=[
        "Acessar a URL raiz da aplicação (ex: https://app.tehillim.com.br).",
        "Aguardar o carregamento completo.",
        "Verificar todos os elementos visuais.",
    ],
    expected=[
        "A landing page é exibida corretamente.",
        "O título/logo da plataforma está visível.",
        "Existe um botão/link para 'Entrar' ou 'Login'.",
        "Nenhum erro JavaScript no console.",
        "A página carrega em menos de 5 segundos.",
    ]
)

add_test_table(doc,
    test_id="NAV-002",
    test_name="Navegação pelo menu lateral — Aluno",
    objective="Verificar que todos os itens do menu lateral do aluno funcionam e redirecionam para as páginas corretas.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Clicar em 'Início' no menu lateral.",
        "Verificar URL e conteúdo da página.",
        "Clicar em 'Trilhas' no menu lateral.",
        "Clicar em 'Aulas' no menu lateral.",
        "Clicar em 'Desempenho' no menu lateral.",
        "Clicar em 'Conquistas' no menu lateral.",
        "Clicar em 'Mensagens' no menu lateral.",
        "Clicar em 'Notificações' no menu lateral.",
        "Clicar em 'Configurações' no menu lateral.",
    ],
    expected=[
        "Cada clique navega para a página correspondente.",
        "A URL muda corretamente a cada navegação.",
        "O item de menu ativo fica visualmente destacado.",
        "Não há erros 404 ou 500 em nenhuma rota.",
        "O conteúdo de cada página é carregado corretamente.",
    ]
)

add_test_table(doc,
    test_id="NAV-003",
    test_name="URL inválida — Página 404",
    objective="Verificar que URLs inexistentes exibem página de erro amigável.",
    preconditions=["Usuário logado."],
    steps=[
        "Na barra de endereço, digitar uma URL inexistente (ex: /pagina-que-nao-existe).",
        "Pressionar Enter.",
    ],
    expected=[
        "Uma página de erro 404 é exibida.",
        "A mensagem informa que a página não foi encontrada.",
        "Existe um link ou botão para retornar à home.",
        "O layout da aplicação (header, navegação) ainda está presente.",
    ]
)

add_test_table(doc,
    test_id="NAV-004",
    test_name="Botão voltar do navegador",
    objective="Verificar que o botão voltar do navegador funciona corretamente.",
    preconditions=["Usuário logado."],
    steps=[
        "Navegar por ao menos 3 páginas diferentes (ex: Home → Trilhas → Desempenho).",
        "Clicar no botão 'Voltar' do navegador.",
        "Clicar novamente em 'Voltar'.",
    ],
    expected=[
        "O navegador retorna à página anterior corretamente.",
        "O conteúdo da página anterior é exibido sem erros.",
        "A sessão permanece ativa.",
    ]
)

add_test_table(doc,
    test_id="NAV-005",
    test_name="Logo / nome da aplicação — Link para home",
    objective="Verificar que clicar no logo ou nome da aplicação redireciona para a home.",
    preconditions=["Usuário logado. Usuário em qualquer página."],
    steps=[
        "Navegar para qualquer página (ex: /desempenho).",
        "Localizar o logo ou nome da aplicação no header/sidebar.",
        "Clicar no logo.",
    ],
    expected=[
        "O usuário é redirecionado para a página inicial (/inicio ou /).",
        "Nenhum erro é exibido.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 3 — DASHBOARD DO ALUNO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 3 — DASHBOARD DO ALUNO (DASH)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a exibição correta do dashboard principal do aluno: XP, streaks, nível, grupos, trilhas e boas-vindas.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="DASH-001",
    test_name="Carregamento do Dashboard — Aluno com dados",
    objective="Verificar que o dashboard carrega corretamente para um aluno com progresso registrado.",
    preconditions=["Usuário logado como aluno.", "Aluno com ao menos 1 módulo iniciado e progresso salvo."],
    steps=[
        "Após o login, acessar /inicio.",
        "Aguardar o carregamento completo da página.",
        "Verificar todos os elementos: header de boas-vindas, estatísticas, grupos.",
    ],
    expected=[
        "A mensagem de boas-vindas exibe o nome do aluno (ex: 'Olá, João!').",
        "O nível atual do aluno é exibido (badge de nível).",
        "O XP total é exibido numericamente.",
        "A barra ou anel de progresso do XP para o próximo nível está visível.",
        "O streak de dias estudados é exibido com o número correto.",
        "Os grupos de conteúdo (Fundamentos do Som, Musicalização, etc.) aparecem como cards.",
        "Cada card de grupo exibe o nome do grupo e uma imagem/ícone.",
        "Nenhum erro JavaScript no console.",
    ]
)

add_test_table(doc,
    test_id="DASH-002",
    test_name="Dashboard — Aluno sem nenhum progresso (conta nova)",
    objective="Verificar que o dashboard exibe estado inicial correto para conta nova sem progresso.",
    preconditions=["Usuário logado como aluno.", "Conta recém criada, sem nenhum progresso."],
    steps=[
        "Acessar /inicio.",
        "Verificar todos os indicadores de progresso.",
    ],
    expected=[
        "XP exibido como 0 (zero) ou nível 1.",
        "Streak exibido como 0 ou 1 (primeiro dia).",
        "Os grupos de conteúdo são listados normalmente.",
        "Não há erros ou tela branca.",
        "A interface exibe estado vazio de forma amigável.",
    ]
)

add_test_table(doc,
    test_id="DASH-003",
    test_name="Cards de Grupos — Clique e navegação",
    objective="Verificar que clicar em um card de grupo navega para a página correta do grupo.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Na home, localizar os cards de grupos (ex: 'Fundamentos do Som').",
        "Clicar no card do primeiro grupo.",
        "Verificar URL e conteúdo.",
        "Voltar e clicar em um segundo grupo.",
    ],
    expected=[
        "Ao clicar em um grupo, a URL muda para /grupos/<slug-do-grupo>.",
        "A página do grupo é carregada com os módulos daquele grupo.",
        "O nome do grupo aparece como título da página.",
        "Nenhum erro é exibido.",
    ]
)

add_test_table(doc,
    test_id="DASH-004",
    test_name="Indicador de Streak — Validação",
    objective="Verificar que o contador de streak reflete corretamente os dias consecutivos de estudo.",
    preconditions=["Usuário logado. Aluno com streak de pelo menos 3 dias."],
    steps=[
        "Acessar a home.",
        "Verificar o número exibido no indicador de streak.",
        "Confirmar com o professor/admin se o número é correto.",
    ],
    expected=[
        "O número de dias no streak corresponde aos dias consecutivos registrados.",
        "O ícone de streak (fogo ou similar) está visível.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 4 — GRUPOS DE CONTEÚDO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 4 — GRUPOS DE CONTEÚDO (GRP)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a listagem e acesso aos grupos de módulos. Cada grupo contém múltiplos módulos de aprendizado.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="GRP-001",
    test_name="Página de Grupo — Carregamento e listagem de módulos",
    objective="Verificar que a página de um grupo exibe corretamente todos os módulos daquele grupo.",
    preconditions=["Usuário logado.", "Grupo com módulos cadastrados."],
    steps=[
        "Navegar para /grupos/fundamentos-do-som (ou outro grupo).",
        "Aguardar o carregamento.",
        "Verificar a lista de módulos exibidos.",
    ],
    expected=[
        "O título do grupo é exibido no topo da página.",
        "Os módulos do grupo são listados como cards ou itens.",
        "Cada módulo exibe: número, título e status de acesso (bloqueado/disponível).",
        "Módulos liberados ao aluno são clicáveis.",
        "Módulos bloqueados exibem indicação visual de bloqueio (cadeado ou similar).",
    ]
)

add_test_table(doc,
    test_id="GRP-002",
    test_name="Acesso a módulo liberado — Clique no card",
    objective="Verificar que clicar em um módulo liberado navega para a página correta do módulo.",
    preconditions=["Usuário logado como aluno.", "Aluno tem acesso ao módulo específico."],
    steps=[
        "Navegar para a página de um grupo.",
        "Localizar um módulo ao qual o aluno tem acesso.",
        "Clicar no módulo.",
    ],
    expected=[
        "A URL muda para /modulos/<slug-do-modulo>.",
        "A página do módulo é carregada corretamente.",
        "O primeiro step disponível é mostrado.",
    ]
)

add_test_table(doc,
    test_id="GRP-003",
    test_name="Módulo bloqueado — Tentativa de acesso direto",
    objective="Verificar que módulos bloqueados não podem ser acessados diretamente pela URL.",
    preconditions=["Usuário logado como aluno.", "Aluno NÃO tem acesso ao módulo específico."],
    steps=[
        "Identificar o slug de um módulo ao qual o aluno não tem acesso.",
        "Digitar manualmente na barra de endereço: /modulos/<slug-bloqueado>.",
        "Pressionar Enter.",
    ],
    expected=[
        "O acesso é negado.",
        "O usuário é redirecionado para a home ou recebe mensagem de acesso negado.",
        "O conteúdo do módulo bloqueado NÃO é exibido.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 5 — CONTEÚDO E STEPS DO MÓDULO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 5 — CONTEÚDO E STEPS DO MÓDULO (MOD)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a estrutura de um módulo: navegação entre steps (Teoria, Vídeo, Visual, Exercício, Jogo), progresso e estado visual.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="MOD-001",
    test_name="Carregamento do Módulo — Estrutura geral",
    objective="Verificar que a página de um módulo carrega com a estrutura completa: sidebar de steps, área de conteúdo.",
    preconditions=["Usuário logado como aluno.", "Módulo com acesso liberado."],
    steps=[
        "Navegar para /modulos/<slug> de um módulo com acesso.",
        "Aguardar o carregamento completo.",
        "Verificar a estrutura visual da página.",
    ],
    expected=[
        "O título do módulo é exibido no topo.",
        "Uma barra/sidebar lateral mostra os steps disponíveis (ex: Teoria, Vídeo, Visual, Exercício, Jogo).",
        "O primeiro step é carregado automaticamente na área principal.",
        "Os steps concluídos estão marcados visualmente (check ou cor diferente).",
        "O step atual está destacado (ativo).",
        "Steps bloqueados estão indicados visualmente.",
    ]
)

add_test_table(doc,
    test_id="MOD-002",
    test_name="Step de Teoria — Exibição do conteúdo textual",
    objective="Verificar que o step de teoria exibe o texto educativo formatado corretamente.",
    preconditions=["Módulo com step de teoria. Usuário com acesso."],
    steps=[
        "Acessar um módulo e navegar para o step 'Teoria'.",
        "Verificar o conteúdo exibido.",
    ],
    expected=[
        "O texto de teoria é exibido completamente.",
        "Formatação do texto (negrito, itálico, listas) está correta.",
        "Imagens dentro da teoria (se houver) são carregadas.",
        "Notações musicais ou exemplos visuais (se houver) estão visíveis.",
        "O botão 'Próximo' ou 'Avançar' está disponível para ir ao próximo step.",
    ]
)

add_test_table(doc,
    test_id="MOD-003",
    test_name="Step de Vídeo — Carregamento do player YouTube",
    objective="Verificar que o vídeo do YouTube incorporado carrega e reproduz corretamente.",
    preconditions=["Módulo com step de vídeo. Conexão com internet."],
    steps=[
        "Acessar o step de vídeo de um módulo.",
        "Aguardar o carregamento do player.",
        "Clicar no botão play para iniciar o vídeo.",
        "Pausar e retomar o vídeo.",
    ],
    expected=[
        "O player do YouTube é carregado dentro da página.",
        "O vídeo reproduz corretamente.",
        "Os controles do player (play, pause, volume) funcionam.",
        "O vídeo ocupa o espaço correto sem quebrar o layout.",
        "Não há erros de carregamento (ex: 'Vídeo não disponível').",
    ]
)

add_test_table(doc,
    test_id="MOD-004",
    test_name="Step Visual — Diagramas e notação musical (VexFlow)",
    objective="Verificar que o step visual exibe imagens, diagramas ou notação musical corretamente.",
    preconditions=["Módulo com step visual."],
    steps=[
        "Acessar o step 'Visual' de um módulo.",
        "Aguardar o carregamento de todos os elementos visuais.",
    ],
    expected=[
        "Imagens e diagramas são exibidos sem distorção.",
        "Notações musicais (VexFlow) são renderizadas corretamente se presentes.",
        "Não há imagens quebradas (ícone de imagem ausente).",
        "O conteúdo é legível e organizado.",
    ]
)

add_test_table(doc,
    test_id="MOD-005",
    test_name="Navegação entre steps — Avançar e voltar",
    objective="Verificar que o aluno pode navegar sequencialmente pelos steps do módulo.",
    preconditions=["Módulo com múltiplos steps. Usuário com acesso."],
    steps=[
        "Acessar um módulo.",
        "Clicar em 'Próximo' ou no botão de avanço para ir ao step 2.",
        "Verificar que o step 2 é carregado.",
        "Continuar avançando até o último step.",
        "Tentar clicar em 'Voltar' para retornar ao step anterior.",
    ],
    expected=[
        "Cada clique em 'Próximo' avança para o step seguinte.",
        "O conteúdo de cada step é carregado corretamente.",
        "A sidebar de steps atualiza o status (ativo, concluído) conforme a navegação.",
        "O botão 'Voltar' retorna ao step anterior sem perder o progresso.",
        "No primeiro step, o botão 'Voltar' está desabilitado ou oculto.",
    ]
)

add_test_table(doc,
    test_id="MOD-006",
    test_name="Barra de progresso do módulo",
    objective="Verificar que o percentual de progresso do módulo é atualizado conforme os steps são concluídos.",
    preconditions=["Módulo com acesso. Progresso parcial ou zerado."],
    steps=[
        "Acessar um módulo.",
        "Notar o percentual de progresso inicial.",
        "Concluir um step.",
        "Verificar o progresso atualizado.",
        "Concluir todos os steps.",
    ],
    expected=[
        "O progresso inicial reflete os steps já concluídos anteriormente.",
        "Ao concluir um step, o percentual de progresso aumenta.",
        "Ao concluir todos os steps, o progresso chega a 100%.",
        "O progresso é salvo mesmo após recarregar a página.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 6 — EXERCÍCIOS
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 6 — EXERCÍCIOS (EXR)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida os diferentes tipos de exercício (múltipla escolha, verdadeiro/falso, preenchimento, etc.) dentro dos módulos.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="EXR-001",
    test_name="Exercício de Múltipla Escolha — Resposta correta",
    objective="Verificar que selecionar a resposta correta em um exercício de múltipla escolha registra o acerto.",
    preconditions=["Módulo com exercício de múltipla escolha. Usuário com acesso."],
    steps=[
        "Acessar o step de exercício de um módulo.",
        "Identificar a pergunta de múltipla escolha.",
        "Selecionar a opção correta.",
        "Confirmar ou enviar a resposta.",
    ],
    expected=[
        "A resposta correta é destacada visualmente (verde ou ícone de acerto).",
        "Uma mensagem de feedback positivo é exibida (ex: 'Correto!', 'Parabéns!').",
        "O progresso do exercício é registrado.",
        "O botão para avançar é habilitado ou aparece.",
    ]
)

add_test_table(doc,
    test_id="EXR-002",
    test_name="Exercício de Múltipla Escolha — Resposta incorreta",
    objective="Verificar que selecionar a resposta errada exibe feedback de erro e indica a correta.",
    preconditions=["Módulo com exercício de múltipla escolha."],
    steps=[
        "Acessar o step de exercício.",
        "Selecionar intencionalmente uma opção errada.",
        "Confirmar a resposta.",
    ],
    expected=[
        "A resposta incorreta é destacada visualmente (vermelho ou ícone de erro).",
        "A resposta correta é indicada visualmente.",
        "Uma mensagem de feedback informativa é exibida (ex: 'Incorreto. A resposta correta é...').",
        "O erro é registrado para histórico de desempenho.",
    ]
)

add_test_table(doc,
    test_id="EXR-003",
    test_name="Exercício Verdadeiro ou Falso",
    objective="Verificar que exercícios de verdadeiro/falso funcionam corretamente.",
    preconditions=["Módulo com exercício verdadeiro/falso."],
    steps=[
        "Acessar step de exercício verdadeiro/falso.",
        "Clicar em 'Verdadeiro'.",
        "Verificar feedback.",
        "Em outro exercício ou nova tentativa, clicar em 'Falso'.",
    ],
    expected=[
        "Ambas as opções (Verdadeiro/Falso) são clicáveis.",
        "Feedback correto é exibido para cada escolha.",
        "Resposta é registrada.",
    ]
)

add_test_table(doc,
    test_id="EXR-004",
    test_name="Exercício de Preenchimento de Lacuna",
    objective="Verificar que exercícios de fill-in-the-blank (preencher espaço) funcionam corretamente.",
    preconditions=["Módulo com exercício de preenchimento."],
    steps=[
        "Acessar step com exercício de preenchimento.",
        "Localizar o campo de texto para preenchimento.",
        "Digitar uma resposta correta.",
        "Confirmar.",
        "Testar com resposta incorreta.",
    ],
    expected=[
        "O campo de texto aceita entrada do teclado.",
        "Resposta correta gera feedback positivo.",
        "Resposta incorreta gera feedback de erro com a resposta certa.",
        "Não é possível avançar sem responder.",
    ]
)

add_test_table(doc,
    test_id="EXR-005",
    test_name="Persistência de respostas de exercício",
    objective="Verificar que as respostas dos exercícios são salvas e visíveis no histórico de desempenho.",
    preconditions=["Aluno completou pelo menos 1 exercício."],
    steps=[
        "Completar um exercício em qualquer módulo.",
        "Navegar para /desempenho.",
        "Verificar se o exercício aparece no histórico.",
    ],
    expected=[
        "O exercício respondido aparece no histórico de respostas.",
        "O status (correto/incorreto) está registrado.",
        "A data/hora está registrada.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 7 — JOGOS INTERATIVOS
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 7 — JOGOS INTERATIVOS (GAME)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a página de jogos (/jogos) e os diferentes tipos de jogo interativo disponíveis na plataforma.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="GAME-001",
    test_name="Página de Jogos — Carregamento e listagem",
    objective="Verificar que a página /jogos carrega e exibe todos os jogos disponíveis organizados por categoria.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Clicar em 'Aulas' no menu lateral (ou navegar diretamente para /jogos).",
        "Aguardar carregamento.",
        "Verificar as categorias de jogos exibidas.",
    ],
    expected=[
        "A página /jogos é carregada sem erros.",
        "Os jogos são listados por categoria (Teoria, Leitura, Ritmo, Som, Interativo, Criativo).",
        "Cada jogo exibe: título, descrição e categoria.",
        "Os jogos são clicáveis.",
        "A barra de navegação por categorias funciona.",
    ]
)

add_test_table(doc,
    test_id="GAME-002",
    test_name="Jogo Verdadeiro ou Falso Rápido",
    objective="Verificar funcionamento do jogo de V/F rápido.",
    preconditions=["Usuário logado. Acesso à página de jogos."],
    steps=[
        "Localizar e abrir o jogo 'Verdadeiro/Falso Rápido'.",
        "Responder algumas afirmações como V ou F.",
        "Completar o jogo.",
    ],
    expected=[
        "O jogo inicia sem erros.",
        "As afirmações são exibidas corretamente.",
        "Os botões V e F respondem ao clique.",
        "Feedback visual é exibido para cada resposta.",
        "Ao finalizar, um resultado/score é mostrado.",
    ]
)

add_test_table(doc,
    test_id="GAME-003",
    test_name="Jogo de Memória Musical",
    objective="Verificar funcionamento do jogo de memória.",
    preconditions=["Usuário logado."],
    steps=[
        "Abrir o jogo 'Memória'.",
        "Clicar em pares de cartas.",
        "Completar o jogo encontrando todos os pares.",
    ],
    expected=[
        "As cartas são exibidas viradas para baixo inicialmente.",
        "Clicar em uma carta a revela.",
        "Par correto: cartas ficam viradas/destacadas.",
        "Par incorreto: cartas voltam a ficar viradas após breve exibição.",
        "Ao completar todos os pares, uma mensagem de conclusão é exibida.",
    ]
)

add_test_table(doc,
    test_id="GAME-004",
    test_name="Jogo Quiz — Fluxo completo",
    objective="Verificar o fluxo completo de um jogo de quiz com pontuação.",
    preconditions=["Usuário logado."],
    steps=[
        "Abrir qualquer jogo de Quiz.",
        "Responder todas as perguntas.",
        "Verificar pontuação final.",
    ],
    expected=[
        "Cada pergunta é exibida claramente.",
        "As opções de resposta são clicáveis.",
        "Feedback é dado após cada resposta.",
        "Ao final, a pontuação total é exibida.",
        "Existe opção para reiniciar ou voltar.",
    ]
)

add_test_table(doc,
    test_id="GAME-005",
    test_name="Jogo dentro de módulo — Step de jogo",
    objective="Verificar que jogos incorporados dentro de módulos funcionam e registram a conclusão.",
    preconditions=["Módulo com step de jogo. Aluno com acesso."],
    steps=[
        "Acessar um módulo que tenha step de jogo.",
        "Navegar até o step de jogo.",
        "Jogar e completar o desafio.",
    ],
    expected=[
        "O jogo do módulo carrega sem erros.",
        "O jogo funciona interativamente.",
        "Ao completar, o step é marcado como concluído.",
        "O progresso do módulo aumenta.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 8 — PLAYER BONA
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 8 — PLAYER BONA — PARTITURA (BONA)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o player de partituras MusicXML disponível nos módulos do Método Bona (módulos 101 a 140).", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="BONA-001",
    test_name="Carregamento do Player de Partitura",
    objective="Verificar que a página de um módulo Bona carrega o player de partitura MusicXML corretamente.",
    preconditions=["Aluno com acesso a um módulo Bona (ex: bona-101). Usuário logado."],
    steps=[
        "Acessar /bona-player/bona-101 (ou equivalente).",
        "Aguardar o carregamento da partitura.",
    ],
    expected=[
        "A partitura musical é renderizada visualmente na tela.",
        "As notas musicais, claves, compassos e dinâmicas são exibidas corretamente.",
        "Não há erros de carregamento do arquivo MusicXML.",
        "O layout da partitura é legível e proporcionado.",
    ]
)

add_test_table(doc,
    test_id="BONA-002",
    test_name="Controles do Player Bona",
    objective="Verificar que os controles de reprodução e navegação da partitura funcionam.",
    preconditions=["Player Bona carregado com partitura."],
    steps=[
        "Clicar no botão Play/Reproduzir.",
        "Clicar em Pause.",
        "Clicar em Stop/Reiniciar.",
        "Ajustar o andamento (BPM) se disponível.",
        "Navegar para uma seção específica da partitura.",
    ],
    expected=[
        "Play inicia a reprodução (se houver playback de metrônomo/MIDI).",
        "Pause interrompe a reprodução.",
        "Stop reinicia para o início.",
        "Ajuste de BPM altera a velocidade.",
        "A partitura destaca a seção em reprodução visualmente.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 9 — TRILHAS DE APRENDIZADO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 9 — TRILHAS DE APRENDIZADO (TRAIL)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a página /trilhas que exibe os caminhos de aprendizado organizados por progressão.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="TRAIL-001",
    test_name="Página de Trilhas — Carregamento",
    objective="Verificar que a página /trilhas exibe as trilhas de aprendizado disponíveis.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Clicar em 'Trilhas' no menu lateral.",
        "Aguardar carregamento da página.",
    ],
    expected=[
        "A página /trilhas é carregada sem erros.",
        "As trilhas de aprendizado são exibidas (ex: trilha de iniciante, intermediário, etc.).",
        "Cada trilha mostra os grupos/módulos pertencentes.",
        "O progresso do aluno em cada trilha é visível.",
    ]
)

add_test_table(doc,
    test_id="TRAIL-002",
    test_name="Navegação da Trilha para Grupo/Módulo",
    objective="Verificar que clicar em um item da trilha navega para o conteúdo correspondente.",
    preconditions=["Página de trilhas carregada."],
    steps=[
        "Localizar uma trilha com módulos listados.",
        "Clicar em um módulo ou grupo dentro da trilha.",
    ],
    expected=[
        "A navegação leva para o grupo ou módulo correto.",
        "A URL muda conforme esperado.",
        "O conteúdo é exibido sem erros.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 10 — DESEMPENHO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 10 — DESEMPENHO (PERF)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a página /desempenho que exibe métricas de progresso, histórico de exercícios e atividade do aluno.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="PERF-001",
    test_name="Página de Desempenho — Carregamento e métricas",
    objective="Verificar que a página /desempenho carrega e exibe as métricas corretas do aluno.",
    preconditions=["Usuário logado como aluno.", "Aluno com algum progresso registrado."],
    steps=[
        "Clicar em 'Desempenho' no menu lateral.",
        "Aguardar carregamento completo.",
        "Verificar todos os indicadores e gráficos.",
    ],
    expected=[
        "A página carrega sem erros.",
        "Os módulos iniciados/concluídos são listados com percentual.",
        "O histórico de exercícios (corretos e incorretos) é exibido.",
        "As datas de atividade são listadas no histórico de streak.",
        "Os dados exibidos correspondem ao histórico real do aluno.",
    ]
)

add_test_table(doc,
    test_id="PERF-002",
    test_name="Desempenho — Aluno sem progresso",
    objective="Verificar estado vazio na página de desempenho para aluno sem histórico.",
    preconditions=["Aluno logado. Sem nenhum progresso ou exercício feito."],
    steps=[
        "Acessar /desempenho.",
        "Verificar o estado exibido.",
    ],
    expected=[
        "A página carrega sem erros.",
        "Uma mensagem de estado vazio é exibida (ex: 'Você ainda não iniciou nenhum módulo').",
        "Não há gráficos ou tabelas com dados incorretos (NaN, undefined, etc.).",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 11 — CONQUISTAS
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 11 — CONQUISTAS (ACH)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a página /conquistas com os badges e realizações do aluno.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="ACH-001",
    test_name="Página de Conquistas — Carregamento",
    objective="Verificar que a página /conquistas exibe as conquistas do aluno.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Clicar em 'Conquistas' no menu lateral.",
        "Aguardar carregamento.",
    ],
    expected=[
        "A página carrega sem erros.",
        "As conquistas desbloqueadas são exibidas com badge/ícone.",
        "As conquistas bloqueadas são exibidas como cinza ou com cadeado.",
        "A descrição de cada conquista é legível.",
    ]
)

add_test_table(doc,
    test_id="ACH-002",
    test_name="Conquista desbloqueada após ação",
    objective="Verificar que completar uma ação que gera conquista a desbloqueia visualmente.",
    preconditions=["Aluno que acabou de completar um módulo pela primeira vez."],
    steps=[
        "Completar um módulo inteiro pela primeira vez.",
        "Navegar para /conquistas.",
        "Verificar se a conquista de 'primeiro módulo' ou equivalente foi desbloqueada.",
    ],
    expected=[
        "A conquista relevante aparece como desbloqueada.",
        "Uma notificação ou feedback visual pode ter sido exibido no momento do desbloqueio.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 12 — MENSAGENS (ALUNO)
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 12 — MENSAGENS — ALUNO (MSG)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o sistema de mensagens entre aluno e professor: envio, recebimento, leitura e badge de não lidas.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="MSG-001",
    test_name="Página de Mensagens — Carregamento",
    objective="Verificar que a página /mensagens exibe o histórico de mensagens do aluno.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Clicar em 'Mensagens' no menu lateral.",
        "Aguardar carregamento.",
    ],
    expected=[
        "A página /mensagens é carregada sem erros.",
        "O histórico de mensagens (enviadas e recebidas) é exibido.",
        "Mensagens não lidas são visualmente diferenciadas (negrito, cor diferente).",
        "Cada mensagem exibe: remetente, conteúdo resumido e data.",
    ]
)

add_test_table(doc,
    test_id="MSG-002",
    test_name="Envio de mensagem pelo aluno",
    objective="Verificar que o aluno consegue enviar uma mensagem ao professor.",
    preconditions=["Usuário logado como aluno."],
    steps=[
        "Acessar /mensagens.",
        "Localizar o campo de composição de nova mensagem.",
        "Digitar um texto (ex: 'Tenho uma dúvida sobre o módulo 5').",
        "Clicar em 'Enviar'.",
    ],
    expected=[
        "A mensagem é enviada com sucesso.",
        "A mensagem aparece no histórico da conversa.",
        "Feedback visual de confirmação (ex: 'Mensagem enviada').",
        "O campo de texto é limpo após o envio.",
    ]
)

add_test_table(doc,
    test_id="MSG-003",
    test_name="Envio de mensagem vazia",
    objective="Verificar que não é possível enviar mensagem sem conteúdo.",
    preconditions=["Usuário logado como aluno. Na página de mensagens."],
    steps=[
        "Deixar o campo de mensagem em branco.",
        "Clicar em 'Enviar'.",
    ],
    expected=[
        "A mensagem NÃO é enviada.",
        "Uma validação é exibida (ex: 'Digite uma mensagem').",
    ]
)

add_test_table(doc,
    test_id="MSG-004",
    test_name="Badge de mensagens não lidas — Atualização",
    objective="Verificar que o badge de não lidas no menu lateral é atualizado quando há mensagens novas.",
    preconditions=["Professor enviou mensagem ao aluno. Aluno ainda não visualizou."],
    steps=[
        "Fazer login como aluno.",
        "Verificar o menu lateral no item 'Mensagens'.",
    ],
    expected=[
        "Um badge com número (ex: '1') aparece ao lado de 'Mensagens'.",
        "O número corresponde à quantidade de mensagens não lidas.",
        "Após abrir /mensagens, o badge desaparece ou diminui.",
    ]
)

add_test_table(doc,
    test_id="MSG-005",
    test_name="Mensagem contextual — Enviada a partir de módulo",
    objective="Verificar que mensagens enviadas dentro de um módulo incluem o contexto do módulo.",
    preconditions=["Módulo com opção de enviar mensagem ao professor. Aluno dentro do módulo."],
    steps=[
        "Dentro de um módulo, localizar botão 'Tirar dúvida' ou 'Enviar mensagem'.",
        "Clicar no botão.",
        "Digitar uma dúvida e enviar.",
        "Verificar a mensagem recebida pelo professor (ou no histórico).",
    ],
    expected=[
        "A mensagem é enviada com contexto: nome do módulo e step atual.",
        "A mensagem aparece no histórico vinculada ao módulo.",
        "O professor (quando verificar) vê qual módulo/step originou a dúvida.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 13 — NOTIFICAÇÕES
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 13 — NOTIFICAÇÕES (NOTIF)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida a página /notificacoes e o sistema de alertas da plataforma.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="NOTIF-001",
    test_name="Página de Notificações — Carregamento",
    objective="Verificar que a página /notificacoes carrega corretamente.",
    preconditions=["Usuário logado."],
    steps=[
        "Clicar em 'Notificações' no menu lateral.",
        "Aguardar carregamento.",
    ],
    expected=[
        "A página /notificacoes é carregada sem erros.",
        "As notificações do usuário são listadas.",
        "Notificações não lidas são visualmente diferenciadas.",
        "Cada notificação exibe: título, descrição e data.",
        "Se não houver notificações, mensagem de estado vazio é exibida.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 14 — CONFIGURAÇÕES E PERFIL
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 14 — CONFIGURAÇÕES E PERFIL (CONF)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida as páginas /configuracoes e /perfil, incluindo a edição do nome de exibição.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="CONF-001",
    test_name="Página de Configurações — Carregamento",
    objective="Verificar que /configuracoes carrega com as opções disponíveis.",
    preconditions=["Usuário logado."],
    steps=[
        "Clicar em 'Configurações' no menu lateral.",
    ],
    expected=[
        "A página carrega sem erros.",
        "As opções de configuração disponíveis são exibidas.",
        "O campo ou seção de nome de exibição está visível.",
    ]
)

add_test_table(doc,
    test_id="CONF-002",
    test_name="Alterar nome de exibição — Sucesso",
    objective="Verificar que o aluno consegue alterar seu nome de exibição na plataforma.",
    preconditions=["Usuário logado. Na página de configurações ou perfil."],
    steps=[
        "Localizar o campo de nome na página de configurações ou perfil.",
        "Limpar o campo atual.",
        "Digitar um novo nome (ex: 'Maria Silva').",
        "Clicar em 'Salvar' ou 'Confirmar'.",
        "Navegar para a home.",
    ],
    expected=[
        "O novo nome é salvo com sucesso.",
        "Uma mensagem de confirmação é exibida (ex: 'Nome atualizado!').",
        "Na home, a saudação (ex: 'Olá, Maria Silva!') exibe o novo nome.",
        "O nome atualizado persiste após recarregar a página.",
    ]
)

add_test_table(doc,
    test_id="CONF-003",
    test_name="Alterar nome — Campo vazio",
    objective="Verificar que não é possível salvar nome vazio.",
    preconditions=["Usuário logado. Na página de configurações."],
    steps=[
        "Limpar completamente o campo de nome.",
        "Clicar em 'Salvar'.",
    ],
    expected=[
        "O formulário NÃO salva.",
        "Validação é exibida (ex: 'O nome não pode estar vazio').",
    ]
)

add_test_table(doc,
    test_id="CONF-004",
    test_name="Página de Perfil — Carregamento",
    objective="Verificar que /perfil exibe as informações do usuário corretamente.",
    preconditions=["Usuário logado."],
    steps=[
        "Navegar para /perfil.",
    ],
    expected=[
        "A página carrega sem erros.",
        "Nome do usuário é exibido.",
        "E-mail do usuário é exibido (ou mascarado).",
        "Informações de nível e XP podem estar presentes.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 15 — GRAVAÇÃO DE ÁUDIO
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 15 — GRAVAÇÃO DE ÁUDIO (REC)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o recurso de gravação de áudio disponível nos módulos, incluindo permissões de microfone, gravação e envio.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="REC-001",
    test_name="Início de gravação — Permissão de microfone",
    objective="Verificar que ao iniciar uma gravação, o navegador solicita permissão de microfone.",
    preconditions=["Módulo com step de gravação. Usuário logado. Microfone disponível."],
    steps=[
        "Acessar um módulo com opção de gravação.",
        "Clicar no botão 'Gravar' ou similar.",
    ],
    expected=[
        "O navegador exibe solicitação de permissão de microfone.",
        "Ao conceder permissão, a gravação inicia.",
        "Um indicador visual (ex: ícone piscando, contador) mostra que está gravando.",
    ]
)

add_test_table(doc,
    test_id="REC-002",
    test_name="Gravação, pausa e envio",
    objective="Verificar o fluxo completo de gravação: iniciar, pausar/parar e enviar.",
    preconditions=["Microfone com permissão. Módulo com gravação."],
    steps=[
        "Iniciar gravação.",
        "Gravar por 5-10 segundos.",
        "Clicar em 'Parar' ou 'Encerrar gravação'.",
        "Verificar o player de pré-visualização do áudio.",
        "Clicar em 'Enviar' ou 'Salvar'.",
    ],
    expected=[
        "A gravação para corretamente.",
        "Um player de áudio com a gravação é exibido para revisão.",
        "O botão de play no player reproduz o áudio gravado.",
        "Ao enviar, uma mensagem de sucesso é exibida.",
        "A gravação aparece no dashboard do professor.",
    ]
)

add_test_table(doc,
    test_id="REC-003",
    test_name="Gravação sem permissão de microfone",
    objective="Verificar o comportamento quando o usuário nega a permissão de microfone.",
    preconditions=["Microfone disponível mas permissão negada."],
    steps=[
        "Tentar iniciar gravação.",
        "Na solicitação do navegador, NEGAR a permissão de microfone.",
    ],
    expected=[
        "Uma mensagem de erro amigável é exibida (ex: 'Por favor, permita o acesso ao microfone').",
        "A gravação NÃO inicia.",
        "A aplicação não trava ou exibe tela branca.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 16 — DASHBOARD DO PROFESSOR
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 16 — DASHBOARD DO PROFESSOR (TDASH)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o dashboard exclusivo do professor em /dashboard, com abas de Alunos, Gravações e Progresso.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="TDASH-001",
    test_name="Carregamento do Dashboard do Professor",
    objective="Verificar que o dashboard do professor carrega corretamente com todas as abas.",
    preconditions=["Usuário logado como PROFESSOR."],
    steps=[
        "Acessar /dashboard.",
        "Aguardar carregamento.",
        "Verificar as abas disponíveis.",
    ],
    expected=[
        "O dashboard carrega sem erros.",
        "Três abas são exibidas: 'Alunos', 'Gravações', 'Progresso'.",
        "A aba padrão ('Alunos') está selecionada e exibe a lista de alunos.",
        "O número total de alunos é exibido.",
    ]
)

add_test_table(doc,
    test_id="TDASH-002",
    test_name="Aba Alunos — Listagem",
    objective="Verificar que a aba Alunos exibe todos os alunos cadastrados com informações básicas.",
    preconditions=["Professor logado. Ao menos 1 aluno cadastrado."],
    steps=[
        "No dashboard, verificar que a aba 'Alunos' está ativa.",
        "Verificar os dados de cada aluno listado.",
    ],
    expected=[
        "Cada aluno é listado com: nome, e-mail, nível, XP, streak.",
        "Existe um link para ver detalhes de cada aluno.",
        "A lista é organizada de forma clara.",
    ]
)

add_test_table(doc,
    test_id="TDASH-003",
    test_name="Aba Gravações — Listagem de submissões",
    objective="Verificar que a aba Gravações exibe as gravações enviadas pelos alunos.",
    preconditions=["Professor logado. Ao menos 1 aluno enviou uma gravação."],
    steps=[
        "No dashboard, clicar na aba 'Gravações'.",
        "Verificar as gravações listadas.",
    ],
    expected=[
        "A aba carrega sem erros.",
        "As gravações aparecem com: nome do aluno, módulo, data de envio.",
        "Existe um player ou link para ouvir cada gravação.",
        "O status da gravação é exibido (pendente, avaliado).",
    ]
)

add_test_table(doc,
    test_id="TDASH-004",
    test_name="Aba Progresso — Visão geral do turma",
    objective="Verificar que a aba Progresso exibe métricas gerais de progresso dos alunos.",
    preconditions=["Professor logado. Alunos com progresso registrado."],
    steps=[
        "No dashboard, clicar na aba 'Progresso'.",
    ],
    expected=[
        "A aba carrega sem erros.",
        "Métricas de progresso são exibidas por aluno ou por módulo.",
        "Os dados refletem o progresso real dos alunos.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 17 — PAINEL ADMIN DO PROFESSOR
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 17 — PAINEL ADMIN DO PROFESSOR (TADM)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o painel administrativo em /admin, onde o professor gerencia contas de alunos.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="TADM-001",
    test_name="Carregamento da Página Admin",
    objective="Verificar que /admin carrega corretamente para o professor.",
    preconditions=["Usuário logado como PROFESSOR."],
    steps=[
        "Clicar em 'Admin' no menu lateral.",
        "Aguardar carregamento de /admin.",
    ],
    expected=[
        "A página admin carrega sem erros.",
        "Um formulário de criação de usuário é exibido.",
        "A lista de alunos existentes é exibida.",
        "Estatísticas gerais (total de alunos, etc.) são visíveis.",
    ]
)

add_test_table(doc,
    test_id="TADM-002",
    test_name="Criar novo aluno — Dados válidos",
    objective="Verificar que o professor consegue criar um novo aluno com dados válidos.",
    preconditions=["Professor logado na página /admin."],
    steps=[
        "Preencher o campo 'Nome' com nome completo.",
        "Preencher o campo 'E-mail' com e-mail válido e único.",
        "Preencher o campo 'Senha' com senha válida (mínimo esperado de caracteres).",
        "Clicar em 'Criar' ou 'Cadastrar Aluno'.",
    ],
    expected=[
        "O aluno é criado com sucesso.",
        "Uma mensagem de confirmação é exibida (ex: 'Aluno criado com sucesso!').",
        "O novo aluno aparece na lista de alunos.",
        "O aluno pode fazer login com as credenciais criadas.",
    ]
)

add_test_table(doc,
    test_id="TADM-003",
    test_name="Criar aluno — E-mail já cadastrado",
    objective="Verificar que o sistema impede a criação de aluno com e-mail duplicado.",
    preconditions=["Professor logado. E-mail de aluno já cadastrado na base."],
    steps=[
        "No formulário de criação, inserir um e-mail já existente.",
        "Preencher nome e senha.",
        "Clicar em 'Criar'.",
    ],
    expected=[
        "O sistema NÃO cria o usuário duplicado.",
        "Mensagem de erro é exibida (ex: 'E-mail já cadastrado' ou 'Usuário já existe').",
    ]
)

add_test_table(doc,
    test_id="TADM-004",
    test_name="Criar aluno — Campos obrigatórios vazios",
    objective="Verificar validação de campos obrigatórios no formulário de criação.",
    preconditions=["Professor logado na página /admin."],
    steps=[
        "Deixar o campo 'Nome' vazio.",
        "Clicar em 'Criar'.",
        "Verificar.",
        "Repetir deixando 'E-mail' vazio.",
        "Repetir deixando 'Senha' vazia.",
    ],
    expected=[
        "Em todos os casos, o formulário NÃO é submetido.",
        "Validação é exibida para o campo faltante.",
    ]
)

add_test_table(doc,
    test_id="TADM-005",
    test_name="Criar aluno — E-mail com formato inválido",
    objective="Verificar que o sistema rejeita e-mails mal formatados.",
    preconditions=["Professor logado na página /admin."],
    steps=[
        "Inserir um e-mail inválido (ex: 'emailsemarroba', 'teste@', '@dominio.com').",
        "Preencher os demais campos.",
        "Clicar em 'Criar'.",
    ],
    expected=[
        "O sistema rejeita o e-mail inválido.",
        "Validação de formato de e-mail é exibida.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 18 — GESTÃO DE ALUNOS PELO PROFESSOR
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 18 — GESTÃO DE ALUNOS PELO PROFESSOR (TSTU)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida as páginas /professor/alunos e /professor/alunos/<id>: listar, visualizar, gerenciar acesso, comentários, mensagens e aulas extras.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="TSTU-001",
    test_name="Lista de Alunos — Carregamento",
    objective="Verificar que /professor/alunos exibe todos os alunos cadastrados.",
    preconditions=["Professor logado."],
    steps=[
        "Clicar em 'Professor/Alunos' no menu lateral.",
        "Aguardar carregamento de /professor/alunos.",
    ],
    expected=[
        "A lista de alunos é exibida.",
        "Cada aluno mostra: nome, e-mail e informações básicas.",
        "Existe um link para acessar os detalhes de cada aluno.",
    ]
)

add_test_table(doc,
    test_id="TSTU-002",
    test_name="Página de Detalhes do Aluno — Carregamento",
    objective="Verificar que a página de detalhes de um aluno exibe informações completas.",
    preconditions=["Professor logado. Ao menos 1 aluno cadastrado."],
    steps=[
        "Na lista de alunos, clicar em um aluno.",
        "Aguardar carregamento de /professor/alunos/<id>.",
    ],
    expected=[
        "A página do aluno carrega sem erros.",
        "Nome, e-mail, nível, XP e streak do aluno são visíveis.",
        "As seções de: Progresso, Comentários, Mensagens e Exercícios estão acessíveis.",
        "O progresso por módulo do aluno é exibido.",
    ]
)

add_test_table(doc,
    test_id="TSTU-003",
    test_name="Liberar acesso a módulo para aluno",
    objective="Verificar que o professor consegue liberar o acesso de um aluno a um módulo específico.",
    preconditions=["Professor na página de detalhes de um aluno.", "Módulo ainda não liberado para o aluno."],
    steps=[
        "Na página do aluno, localizar a seção de gerenciamento de acesso a módulos.",
        "Localizar um módulo bloqueado.",
        "Ativar o toggle ou checkbox para liberar o módulo.",
        "Salvar ou confirmar.",
        "Fazer login como aluno e verificar se o módulo está acessível.",
    ],
    expected=[
        "O módulo é liberado com sucesso.",
        "Feedback visual de confirmação é exibido (ex: 'Acesso atualizado').",
        "O aluno agora consegue acessar o módulo liberado.",
        "O professor vê o módulo como 'liberado' na interface.",
    ]
)

add_test_table(doc,
    test_id="TSTU-004",
    test_name="Revogar acesso a módulo de aluno",
    objective="Verificar que o professor consegue revogar o acesso de um aluno a um módulo.",
    preconditions=["Professor na página do aluno.", "Módulo liberado para o aluno."],
    steps=[
        "Localizar um módulo liberado.",
        "Desativar o toggle/checkbox para revogar.",
        "Salvar ou confirmar.",
        "Logar como aluno e tentar acessar o módulo revogado.",
    ],
    expected=[
        "O acesso é revogado com sucesso.",
        "O aluno não consegue mais acessar o módulo.",
        "Ao tentar acessar diretamente pela URL, o aluno é redirecionado.",
    ]
)

add_test_table(doc,
    test_id="TSTU-005",
    test_name="Adicionar comentário sobre aluno",
    objective="Verificar que o professor consegue adicionar um comentário/feedback no perfil do aluno.",
    preconditions=["Professor na página de detalhes do aluno."],
    steps=[
        "Localizar a seção 'Comentários'.",
        "Clicar em 'Adicionar comentário'.",
        "Digitar um feedback (ex: 'Aluno demonstra boa evolução no ritmo').",
        "Clicar em 'Salvar'.",
    ],
    expected=[
        "O comentário é salvo com sucesso.",
        "O comentário aparece na lista com: texto, data e professor.",
        "Mensagem de confirmação é exibida.",
    ]
)

add_test_table(doc,
    test_id="TSTU-006",
    test_name="Deletar comentário do aluno",
    objective="Verificar que o professor consegue excluir um comentário.",
    preconditions=["Ao menos 1 comentário cadastrado no aluno."],
    steps=[
        "Localizar um comentário na seção de comentários.",
        "Clicar em 'Excluir' ou ícone de deletar.",
        "Confirmar a exclusão (se houver modal de confirmação).",
    ],
    expected=[
        "O comentário é removido da lista.",
        "Não há mais erros ou rastros do comentário excluído.",
        "Mensagem de confirmação pode ser exibida.",
    ]
)

add_test_table(doc,
    test_id="TSTU-007",
    test_name="Adicionar aula extra para aluno",
    objective="Verificar que o professor consegue adicionar uma aula extra (vídeo, PDF ou link) para um aluno.",
    preconditions=["Professor na página de detalhes do aluno."],
    steps=[
        "Localizar a seção de 'Aulas Extras'.",
        "Clicar em 'Adicionar aula extra'.",
        "Preencher: título, descrição, tipo (vídeo/PDF/link) e URL.",
        "Clicar em 'Salvar'.",
        "Logar como aluno e verificar a aula extra.",
    ],
    expected=[
        "A aula extra é adicionada com sucesso.",
        "O aluno vê a nova aula extra na sua área.",
        "O tipo correto (vídeo/PDF/link) é exibido com o ícone correspondente.",
        "O link/URL funciona ao clicar.",
    ]
)

add_test_table(doc,
    test_id="TSTU-008",
    test_name="Excluir aula extra",
    objective="Verificar que o professor consegue excluir uma aula extra.",
    preconditions=["Ao menos 1 aula extra cadastrada para o aluno."],
    steps=[
        "Localizar a aula extra.",
        "Clicar em 'Excluir'.",
        "Confirmar.",
    ],
    expected=[
        "A aula extra é removida.",
        "O aluno não vê mais a aula extra em sua área.",
    ]
)

add_test_table(doc,
    test_id="TSTU-009",
    test_name="Envio de mensagem pelo professor ao aluno",
    objective="Verificar que o professor consegue enviar mensagem para um aluno específico.",
    preconditions=["Professor na página de detalhes do aluno."],
    steps=[
        "Localizar a seção 'Mensagens'.",
        "Digitar uma mensagem (ex: 'Parabéns pelo progresso!').",
        "Clicar em 'Enviar'.",
        "Logar como aluno e verificar se a mensagem chegou.",
    ],
    expected=[
        "A mensagem é enviada com sucesso.",
        "O aluno recebe a mensagem em /mensagens.",
        "O badge de não lidas é atualizado para o aluno.",
    ]
)

add_test_table(doc,
    test_id="TSTU-010",
    test_name="Visualizar histórico de exercícios do aluno",
    objective="Verificar que o professor consegue ver as respostas dos exercícios do aluno.",
    preconditions=["Aluno com exercícios respondidos.", "Professor na página do aluno."],
    steps=[
        "Localizar a seção 'Exercícios' na página do aluno.",
        "Verificar o histórico de respostas.",
    ],
    expected=[
        "O histórico de exercícios é exibido: pergunta, resposta do aluno, correto/incorreto, data.",
        "As respostas são exibidas de forma clara.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 19 — RESPONSIVIDADE
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 19 — RESPONSIVIDADE (RESP)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o comportamento da interface em diferentes tamanhos de tela: Desktop, Notebook, Tablet e Mobile.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="RESP-001",
    test_name="Layout Desktop (1920x1080 e 1440x900)",
    objective="Verificar que a interface se apresenta corretamente em resoluções de desktop.",
    preconditions=["Usuário logado. Navegador em tela cheia."],
    steps=[
        "Definir o zoom do navegador para 100%.",
        "Acessar as páginas: Home, Módulo, Dashboard Professor, Admin.",
        "Verificar o layout de cada página.",
    ],
    expected=[
        "O menu lateral (sidebar) de 210px é exibido corretamente à esquerda.",
        "O conteúdo principal ocupa o restante da tela.",
        "Não há elementos sobrepostos ou cortados.",
        "Textos são legíveis.",
        "Imagens e ícones estão proporcionados.",
    ]
)

add_test_table(doc,
    test_id="RESP-002",
    test_name="Layout Mobile (375x812 — iPhone)",
    objective="Verificar que a interface se adapta corretamente para telas mobile.",
    preconditions=["Usar DevTools do navegador → emular dispositivo mobile (ex: iPhone 14 Pro)."],
    steps=[
        "Abrir as DevTools e selecionar emulação de iPhone 14 Pro.",
        "Navegar para: Home, Módulo, Mensagens, Configurações.",
        "Verificar cada página.",
    ],
    expected=[
        "O menu lateral DESAPARECE e é substituído por barra de navegação INFERIOR.",
        "A barra inferior tem ícones clicáveis para as principais seções.",
        "Todo o conteúdo cabe na tela sem scroll horizontal.",
        "Textos são legíveis sem zoom.",
        "Botões e links têm área de toque adequada (mínimo 44x44px).",
        "Formulários são usáveis com teclado virtual.",
    ]
)

add_test_table(doc,
    test_id="RESP-003",
    test_name="Layout Tablet (768x1024 — iPad)",
    objective="Verificar que a interface se adapta adequadamente para tablets.",
    preconditions=["Emular iPad via DevTools."],
    steps=[
        "Definir emulação para iPad (768x1024).",
        "Navegar pelas principais páginas.",
    ],
    expected=[
        "O layout se adapta corretamente.",
        "Sidebar pode estar visível ou substituída por hamburger menu.",
        "Conteúdo é legível e bem organizado.",
        "Não há overflow horizontal.",
    ]
)

add_test_table(doc,
    test_id="RESP-004",
    test_name="Player de Módulo em Mobile",
    objective="Verificar que o player de módulo (steps, conteúdo, vídeo) funciona em mobile.",
    preconditions=["Emulação mobile. Aluno com módulo liberado."],
    steps=[
        "Em modo mobile, acessar um módulo.",
        "Navegar entre os steps.",
        "Verificar step de vídeo.",
        "Verificar step de exercício.",
    ],
    expected=[
        "A sidebar de steps se adapta (pode virar abas horizontais ou dropdown).",
        "O conteúdo é legível em 375px.",
        "Vídeo YouTube adapta largura.",
        "Exercícios são interativos por touch.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 20 — SEGURANÇA FUNCIONAL
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 20 — SEGURANÇA FUNCIONAL (SEC)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Testes de segurança do ponto de vista funcional (sem acesso ao código). Valida proteção de rotas, sessões e permissões.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="SEC-001",
    test_name="Acesso direto a API sem autenticação",
    objective="Verificar que endpoints da API requerem autenticação.",
    preconditions=["Nenhuma sessão ativa."],
    steps=[
        "Sem estar logado, tentar acessar diretamente: /api/progress via URL do navegador.",
        "Tentar acessar: /api/messages.",
        "Verificar a resposta.",
    ],
    expected=[
        "Retorno de erro 401 (Unauthorized) ou redirecionamento para login.",
        "Nenhum dado de usuário é retornado.",
    ]
)

add_test_table(doc,
    test_id="SEC-002",
    test_name="Aluno tentando acessar dados de outro aluno via URL",
    objective="Verificar que um aluno não consegue ver dados de outro aluno manipulando a URL.",
    preconditions=["Dois alunos cadastrados: Aluno A e Aluno B."],
    steps=[
        "Logar como Aluno A.",
        "Tentar acessar /professor/alunos/<id-do-aluno-B>.",
        "Tentar acessar /api/teacher/student/<id-do-aluno-B>/access.",
    ],
    expected=[
        "O Aluno A é redirecionado ou recebe 403 (Forbidden).",
        "Os dados do Aluno B não são exibidos.",
    ]
)

add_test_table(doc,
    test_id="SEC-003",
    test_name="Manipulação de URL para módulo bloqueado",
    objective="Verificar que inserir manualmente a URL de um módulo bloqueado não expõe o conteúdo.",
    preconditions=["Aluno logado sem acesso ao módulo X."],
    steps=[
        "Identificar o slug de um módulo ao qual o aluno não tem acesso.",
        "Digitar manualmente /modulos/<slug-bloqueado> na barra de endereço.",
    ],
    expected=[
        "O acesso é negado.",
        "O conteúdo do módulo não é exibido.",
        "Redirecionamento ou mensagem de acesso negado.",
    ]
)

add_test_table(doc,
    test_id="SEC-004",
    test_name="Sessão expirada — Comportamento",
    objective="Verificar que uma sessão expirada redireciona o usuário ao login.",
    preconditions=["Usuário logado. Aguardar expiração da sessão (ou limpar cookies manualmente)."],
    steps=[
        "Logar na aplicação.",
        "Limpar os cookies de sessão via DevTools do navegador.",
        "Tentar navegar para /inicio.",
    ],
    expected=[
        "O usuário é redirecionado para a página de login.",
        "Nenhum dado privado é exibido.",
        "Mensagem informando que a sessão expirou pode ser exibida.",
    ]
)

add_test_table(doc,
    test_id="SEC-005",
    test_name="XSS básico em campos de texto",
    objective="Verificar que campos de texto não executam scripts injetados.",
    preconditions=["Usuário logado. Acesso a formulário de mensagem ou configurações."],
    steps=[
        "Em um campo de texto (ex: campo de mensagem ou nome), inserir: <script>alert('xss')</script>",
        "Enviar o formulário.",
        "Verificar se o alerta é executado.",
    ],
    expected=[
        "Nenhum alerta JavaScript é executado.",
        "O texto é exibido como texto literal (escapado corretamente).",
        "Nenhum código JavaScript é injetado na página.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 21 — PERFORMANCE VISUAL
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 21 — PERFORMANCE VISUAL (VIS)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida tempo de carregamento, ausência de travamentos, componentes quebrados e duplicações.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="VIS-001",
    test_name="Tempo de carregamento das páginas principais",
    objective="Verificar que as páginas carregam em tempo aceitável (< 5 segundos).",
    preconditions=["Conexão com internet (velocidade normal). Usuário logado."],
    steps=[
        "Abrir as DevTools (F12) → aba 'Network'.",
        "Navegar para /inicio e medir o tempo de carregamento.",
        "Repetir para /modulos/<slug>, /desempenho, /dashboard (professor).",
    ],
    expected=[
        "Todas as páginas carregam em menos de 5 segundos em conexão normal.",
        "Não há requisições com timeout ou erro 503.",
        "A interface fica utilizável antes do carregamento completo (loading progressivo).",
    ]
)

add_test_table(doc,
    test_id="VIS-002",
    test_name="Ausência de tela branca",
    objective="Verificar que nenhuma página exibe tela branca (WSOD) após carregamento.",
    preconditions=["Usuário logado."],
    steps=[
        "Navegar para todas as páginas principais.",
        "Forçar recarga (F5 ou Ctrl+R) em cada página.",
    ],
    expected=[
        "Nenhuma página exibe tela branca após carregamento.",
        "Se houver delay, um spinner ou loading skeleton é exibido.",
    ]
)

add_test_table(doc,
    test_id="VIS-003",
    test_name="Ausência de elementos duplicados",
    objective="Verificar que nenhuma página exibe componentes duplicados ou sobrepostos.",
    preconditions=["Usuário logado."],
    steps=[
        "Navegar pelo app e verificar visualmente cada página.",
        "Usar botão 'Voltar' e 'Avançar' do navegador repetidamente.",
        "Verificar se algum componente aparece duplicado.",
    ],
    expected=[
        "Nenhum elemento (menu, header, card) aparece mais de uma vez na mesma página.",
        "Não há sobreposição de modais ou drawers.",
    ]
)

add_test_table(doc,
    test_id="VIS-004",
    test_name="Console sem erros críticos",
    objective="Verificar que o console do navegador não exibe erros JavaScript críticos.",
    preconditions=["DevTools aberto na aba 'Console'."],
    steps=[
        "Navegar por todas as páginas principais com o console aberto.",
        "Observar mensagens de erro (vermelho) no console.",
    ],
    expected=[
        "Não há erros críticos (vermelho) no console.",
        "Warnings (amarelos) são aceitáveis mas devem ser documentados.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 22 — ESTADOS EXTREMOS
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 22 — ESTADOS EXTREMOS (EXTR)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida o comportamento da aplicação em cenários extremos: sem internet, API fora, dados nulos, muitos registros.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="EXTR-001",
    test_name="Comportamento sem conexão com internet",
    objective="Verificar como a aplicação se comporta quando a internet é desconectada.",
    preconditions=["Usuário logado. Conexão com internet disponível."],
    steps=[
        "Abrir a aplicação normalmente.",
        "Desativar a conexão de rede (modo avião ou desligar Wi-Fi).",
        "Tentar navegar para outra página.",
        "Tentar enviar uma mensagem.",
    ],
    expected=[
        "A aplicação exibe uma mensagem de erro de conectividade (ex: 'Sem conexão com internet').",
        "A aplicação NÃO trava ou exibe tela branca.",
        "Ao reconectar, a aplicação funciona normalmente.",
    ]
)

add_test_table(doc,
    test_id="EXTR-002",
    test_name="Aluno sem nenhum módulo liberado",
    objective="Verificar a interface quando um aluno não tem nenhum módulo com acesso.",
    preconditions=["Aluno recém criado sem nenhum módulo liberado."],
    steps=[
        "Logar como aluno sem módulos.",
        "Verificar a home.",
        "Clicar nos grupos.",
    ],
    expected=[
        "A home carrega sem erros.",
        "Os grupos são exibidos.",
        "Na página de um grupo, os módulos aparecem como bloqueados.",
        "Mensagem de orientação é exibida (ex: 'Peça ao seu professor para liberar módulos').",
        "Não há erros de JavaScript ou dados indefinidos.",
    ]
)

add_test_table(doc,
    test_id="EXTR-003",
    test_name="Professor sem alunos cadastrados",
    objective="Verificar a interface do professor quando não há alunos cadastrados.",
    preconditions=["Conta de professor sem alunos."],
    steps=[
        "Logar como professor sem alunos.",
        "Verificar /dashboard e /professor/alunos.",
    ],
    expected=[
        "As páginas carregam sem erros.",
        "Uma mensagem de estado vazio é exibida (ex: 'Nenhum aluno cadastrado ainda').",
        "Não há erros de loop ou crash.",
    ]
)

add_test_table(doc,
    test_id="EXTR-004",
    test_name="Nome de usuário muito longo",
    objective="Verificar o comportamento da interface com nome de usuário extremamente longo.",
    preconditions=["Usuário logado. Na página de configurações."],
    steps=[
        "Alterar o nome para uma string de 200 caracteres.",
        "Salvar.",
        "Verificar o nome exibido na home e no header.",
    ],
    expected=[
        "O nome é truncado ou quebrado corretamente no layout.",
        "A interface não é quebrada ou distorcida.",
        "O nome é salvo (pode ser truncado na exibição).",
    ]
)

add_test_table(doc,
    test_id="EXTR-005",
    test_name="Múltiplos cliques rápidos em botão",
    objective="Verificar que clicar repetidamente em botões de ação não causa duplicação de requisições.",
    preconditions=["Usuário logado."],
    steps=[
        "Em um formulário (ex: envio de mensagem), preencher o campo.",
        "Clicar no botão 'Enviar' rapidamente 5 vezes.",
    ],
    expected=[
        "A mensagem é enviada apenas UMA vez.",
        "O botão fica desabilitado após o primeiro clique.",
        "Não há duplicação de registros.",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# MÓDULO 23 — ACESSIBILIDADE
# ══════════════════════════════════════════════

add_heading(doc, "MÓDULO 23 — ACESSIBILIDADE (ACESS)", level=1, color=(0xC4, 0x94, 0x3A))
add_para(doc, "Valida aspectos básicos de acessibilidade: contraste, navegação por teclado, labels e legibilidade.", italic=True)
doc.add_paragraph()

add_test_table(doc,
    test_id="ACESS-001",
    test_name="Contraste de textos e fundos",
    objective="Verificar que os textos têm contraste suficiente para leitura em diferentes seções.",
    preconditions=["Usuário logado."],
    steps=[
        "Verificar visualmente o contraste em: textos de módulos, botões, labels, mensagens de erro.",
        "Prestar atenção especial ao texto dourado/âmbar sobre fundo branco.",
    ],
    expected=[
        "Todos os textos são legíveis sem esforço.",
        "Textos sobre fundo colorido (ex: botões dourados) são contrastantes.",
        "Mensagens de erro (vermelhas) são legíveis.",
    ]
)

add_test_table(doc,
    test_id="ACESS-002",
    test_name="Navegação por teclado — Formulários",
    objective="Verificar que formulários são navegáveis apenas pelo teclado.",
    preconditions=["Usuário na página de login."],
    steps=[
        "Clicar no campo de e-mail.",
        "Usar TAB para mover para o campo de senha.",
        "Usar TAB para mover para o botão 'Entrar'.",
        "Pressionar ENTER para submeter.",
    ],
    expected=[
        "A tecla TAB move o foco corretamente entre os campos.",
        "O foco é visível (destaque visual no campo ativo).",
        "ENTER submete o formulário.",
    ]
)

add_test_table(doc,
    test_id="ACESS-003",
    test_name="Labels e placeholders em campos",
    objective="Verificar que campos de formulário têm labels ou placeholders claros.",
    preconditions=["Página com formulários (login, criação de usuário, mensagens)."],
    steps=[
        "Verificar cada campo dos formulários.",
        "Verificar se existe label ou placeholder descritivo.",
    ],
    expected=[
        "Cada campo tem label ou placeholder que descreve o que deve ser preenchido.",
        "Campos obrigatórios são indicados visualmente (ex: asterisco *).",
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# RESUMO EXECUTIVO / CONSOLIDAÇÃO
# ══════════════════════════════════════════════

add_heading(doc, "CONSOLIDAÇÃO DOS RESULTADOS", level=1, color=(0x1A, 0x1A, 0x2E))

summary_text = """Após a execução de todos os testes, preencher a tabela abaixo com o resumo dos resultados obtidos."""
add_para(doc, summary_text)
doc.add_paragraph()

# Tabela de resumo por módulo
summary_table = doc.add_table(rows=0, cols=6)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Cabeçalho
header_row = summary_table.add_row()
headers = ["Módulo", "Total de Testes", "Aprovados", "Reprovados", "Bloqueados", "% Aprovação"]
for i, h in enumerate(headers):
    set_cell_bg(header_row.cells[i], "C4943A")
    p = header_row.cells[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

modulos_resumo = [
    ("AUTH — Autenticação", "10", "", "", "", ""),
    ("NAV — Navegação e Rotas", "5", "", "", "", ""),
    ("DASH — Dashboard Aluno", "4", "", "", "", ""),
    ("GRP — Grupos de Conteúdo", "3", "", "", "", ""),
    ("MOD — Módulos e Steps", "6", "", "", "", ""),
    ("EXR — Exercícios", "5", "", "", "", ""),
    ("GAME — Jogos Interativos", "5", "", "", "", ""),
    ("BONA — Player Partitura", "2", "", "", "", ""),
    ("TRAIL — Trilhas", "2", "", "", "", ""),
    ("PERF — Desempenho", "2", "", "", "", ""),
    ("ACH — Conquistas", "2", "", "", "", ""),
    ("MSG — Mensagens Aluno", "5", "", "", "", ""),
    ("NOTIF — Notificações", "1", "", "", "", ""),
    ("CONF — Configurações", "4", "", "", "", ""),
    ("REC — Gravação de Áudio", "3", "", "", "", ""),
    ("TDASH — Dashboard Professor", "4", "", "", "", ""),
    ("TADM — Admin Professor", "5", "", "", "", ""),
    ("TSTU — Gestão de Alunos", "10", "", "", "", ""),
    ("RESP — Responsividade", "4", "", "", "", ""),
    ("SEC — Segurança Funcional", "5", "", "", "", ""),
    ("VIS — Performance Visual", "4", "", "", "", ""),
    ("EXTR — Estados Extremos", "5", "", "", "", ""),
    ("ACESS — Acessibilidade", "3", "", "", "", ""),
    ("TOTAL GERAL", "99", "", "", "", ""),
]

for i, (mod, total, aprov, reprov, bloq, perc) in enumerate(modulos_resumo):
    row = summary_table.add_row()
    values = [mod, total, aprov, reprov, bloq, perc]
    bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
    if i == len(modulos_resumo) - 1:
        bg = "E8C77A"
    for j, val in enumerate(values):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == len(modulos_resumo) - 1:
            run.bold = True

doc.add_paragraph()

# Área de assinaturas
add_heading(doc, "ASSINATURAS", level=2)
doc.add_paragraph()

sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = 'Table Grid'

sig_fields = [
    ("Responsável pelo QA:", "___________________________"),
    ("Data de Conclusão:", "____/____/________"),
    ("Aprovado para Release:", "[ ] SIM   [ ] NÃO"),
]

for i, (label, val) in enumerate(sig_fields):
    set_cell_bg(sig_table.rows[i].cells[0], "E8C77A")
    sig_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
    sig_table.rows[i].cells[1].paragraphs[0].add_run(val)

doc.add_paragraph()
add_para(doc, "Observações finais sobre o release:", bold=True)
for _ in range(5):
    doc.add_paragraph("_______________________________________________")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fim do Documento de QA — Tehillim Music Learning v1.0 —")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ─── SALVAR ─────────────────────────────────
output_path = "QA_Tehillim_Music_Learning_v1.0.docx"
doc.save(output_path)
print(f"\n✅ Documento gerado com sucesso: {output_path}")
print(f"   Tamanho: {__import__('os').path.getsize(output_path) / 1024:.1f} KB")
